import shutil
import os
import yaml
from docx import Document
from copy import deepcopy

def find_and_replace(doc, prefix, replacements):
    # Replace in paragraphs
    for para in doc.paragraphs:
        if prefix in para.text:
            inline = para.runs
            for run in inline:
                if prefix in run.text:
                    for key,value in replacements.items():
                        search_text = prefix+key
                        if search_text in run.text:
                            run.text = run.text.replace(search_text, value)
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if prefix in para.text:
                        for run in para.runs:
                            if prefix in run.text:
                                for key,value in replacements.items():
                                    search_text = prefix+key
                                    if search_text in run.text:
                                        run.text = run.text.replace(search_text, value)

def find_table(doc, prefix):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if prefix in para.text:
                        return table
    return None

def fill_row(prefix, values, row):
    for cell in row.cells:
        for para in cell.paragraphs:
            if prefix in para.text:
                for run in para.runs:
                    if prefix in run.text:
                        for key,value in values.items():
                            search_text = prefix+key
                            if search_text in run.text:
                                run.text = run.text.replace(search_text, str(value))

def copy_paragraphs(source_cell, target_cell):
    # Clear target cell first
    target_cell._element.clear_content()
    for para in source_cell.paragraphs:
        new_para = target_cell.add_paragraph()
        # Copy runs with formatting
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.color.rgb = run.font.color.rgb

def copy_row(table, row_idx):
    tbl = table._tbl
    row = table.rows[row_idx]._tr
    new_row = deepcopy(row)
    tbl.append(new_row)

def fill_table(doc, prefix, values):
    table = find_table(doc, prefix)

    if table == None:
        return

    count = len(values) - 1
    while count != 0:
        copy_row(table, -1)
        count -= 1
    
    count = -1
    while abs(count) <= len(values):
        fill_row(prefix, values[count], table.rows[count])
        count -= 1

def ucarkagen(input, nusga, output):
    """Üç arka Generator: input - yaml file, nusga - marked docx file, output - file to save"""
    # Validate input
    if not os.path.exists(input):
        return

    # Validate nusga
    if not os.path.exists(nusga) or not nusga.lower().endswith(".docx"):
        return
    
    # Validate output
    if not output:
        return

    # Open input
    with open(input, "r") as file:
        data = yaml.safe_load(file)

    # Copy the template
    output = shutil.copy(nusga, output)

    # Open docx
    doc = Document(output)

    if 'men' in data.keys():
        men = data['men']
    
    if 'is-yerler' in men.keys():
        isyerler = men['is-yerler']
        del men['is-yerler']
        fill_table(doc, "%men-is-yerler-", isyerler)

    find_and_replace(doc, '%men-', men)

    if 'garyndas' in data.keys():
        garyndas = data['garyndas']
        fill_table(doc, "%garyndas-", garyndas)

    doc.save(output)
